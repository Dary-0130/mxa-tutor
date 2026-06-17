"""docx parser adapter backed by python-docx."""

from __future__ import annotations

import zipfile
from collections.abc import Iterable
from datetime import UTC, datetime
from pathlib import Path
from typing import Protocol

from docx import Document
from docx.opc.exceptions import PackageNotFoundError

from core.domain.exceptions import DocumentParseError
from core.interfaces.document_parser import (
    DocumentParser,
    FigurePlaceholder,
    ParsedDocument,
    ParsedLocatorIndex,
    compute_file_hash,
)

DOCX_MAGIC = b"PK\x03\x04"
CONTENT_TYPES = "[Content_Types].xml"
DEFAULT_MAX_DOCX_BYTES = 50 * 1024 * 1024
DEFAULT_MAX_UNCOMPRESSED_BYTES = 50 * 1024 * 1024


class _ParagraphLike(Protocol):
    text: str


class DocxParser(DocumentParser):
    """Extract paragraphs, table placeholders, and image placeholders from docx."""

    def __init__(
        self,
        max_file_bytes: int = DEFAULT_MAX_DOCX_BYTES,
        max_uncompressed_bytes: int = DEFAULT_MAX_UNCOMPRESSED_BYTES,
    ) -> None:
        self._max_file_bytes = max_file_bytes
        self._max_uncompressed_bytes = max_uncompressed_bytes

    def supports(self, file_path: Path) -> bool:
        if file_path.suffix.lower() != ".docx" or _read_prefix(file_path, 4) != DOCX_MAGIC:
            return False
        try:
            with zipfile.ZipFile(file_path) as package:
                return CONTENT_TYPES in package.namelist()
        except (OSError, zipfile.BadZipFile):
            return False

    def parse(self, file_path: Path, timeout_seconds: float = 30.0) -> ParsedDocument:
        _ = timeout_seconds
        if file_path.suffix.lower() != ".docx" or _read_prefix(file_path, 4) != DOCX_MAGIC:
            raise DocumentParseError("unsupported_docx_format") from None
        if file_path.stat().st_size > self._max_file_bytes:
            raise DocumentParseError("docx_too_large") from None
        _scan_docx_package(file_path, self._max_uncompressed_bytes)

        try:
            document = Document(str(file_path))
        except (OSError, PackageNotFoundError, ValueError, KeyError, zipfile.BadZipFile):
            raise DocumentParseError("docx_parse_failed") from None

        section_ids: list[str] = []
        text_blocks: list[str] = []
        for index, paragraph in enumerate(_nonempty_paragraphs(document.paragraphs), start=1):
            section_id = f"S{index}"
            section_ids.append(section_id)
            text_blocks.append(f"[{section_id}]\n{paragraph}")

        table_placeholders = [
            f"TABLE-{index:02d}" for index, _ in enumerate(document.tables, start=1)
        ]
        raw_text = "\n\n".join(text_blocks).strip()
        if not raw_text and not table_placeholders and not document.inline_shapes:
            raise DocumentParseError("document_text_too_short") from None

        figures = [
            FigurePlaceholder(f"FIG-{index:02d}", "", None)
            for index, _ in enumerate(document.inline_shapes, start=1)
        ]
        return ParsedDocument(
            raw_text=raw_text,
            page_count=None,
            figure_placeholders=figures,
            table_placeholders=table_placeholders,
            locator_index=ParsedLocatorIndex(
                section_ids=section_ids,
                equation_ids=_extract_equation_ids(raw_text),
                figure_ids=[figure.figure_id for figure in figures],
            ),
            file_hash=compute_file_hash(file_path),
            extracted_at=datetime.now(UTC),
        )


def _read_prefix(file_path: Path, size: int) -> bytes:
    try:
        with file_path.open("rb") as stream:
            return stream.read(size)
    except OSError:
        return b""


def _scan_docx_package(file_path: Path, max_uncompressed_bytes: int) -> None:
    try:
        with zipfile.ZipFile(file_path) as package:
            infos = package.infolist()
            names = {info.filename for info in infos}
            if CONTENT_TYPES not in names or "word/document.xml" not in names:
                raise DocumentParseError("docx_package_invalid") from None
            if any(name.endswith("vbaProject.bin") for name in names):
                raise DocumentParseError("docx_macro_not_supported") from None
            if sum(info.file_size for info in infos) > max_uncompressed_bytes:
                raise DocumentParseError("docx_uncompressed_too_large") from None
    except zipfile.BadZipFile:
        raise DocumentParseError("docx_parse_failed") from None


def _nonempty_paragraphs(paragraphs: Iterable[_ParagraphLike]) -> list[str]:
    return [paragraph.text.strip() for paragraph in paragraphs if paragraph.text.strip()]


def _extract_equation_ids(raw_text: str) -> list[str]:
    lines = [line for line in raw_text.splitlines() if "=" in line]
    return [f"EQ-{index:02d}" for index, _ in enumerate(lines, start=1)]
