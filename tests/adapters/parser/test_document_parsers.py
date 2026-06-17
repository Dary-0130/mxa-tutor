from pathlib import Path

import pytest
from docx import Document

from adapters.parser.docx_parser import DocxParser
from adapters.parser.pdf_parser import PdfParser
from core.domain.exceptions import DocumentParseError
from tests.fixtures.malicious_documents.build_fixtures import build_all


@pytest.fixture(scope="session")
def malicious_document_dir(tmp_path_factory: pytest.TempPathFactory) -> dict[str, Path]:
    return build_all(tmp_path_factory.mktemp("malicious_documents"))


def test_docx_parser_extracts_text_and_table_placeholders(tmp_path: Path) -> None:
    path = tmp_path / "paper.docx"
    document = Document()
    document.add_paragraph("Short circuit experiment report")
    table = document.add_table(rows=1, cols=1)
    table.cell(0, 0).text = "H = 3.5 s"
    document.save(path)

    parsed = DocxParser().parse(path)

    assert parsed.raw_text == "[S1]\nShort circuit experiment report"
    assert parsed.table_placeholders == ["TABLE-01"]
    assert parsed.locator_index.section_ids == ["S1"]


def test_docx_support_requires_extension_magic_and_content_types(tmp_path: Path) -> None:
    wrong_extension = tmp_path / "paper.zip"
    wrong_extension.write_bytes(b"PK\x03\x04")
    missing_content_types = tmp_path / "paper.docx"
    missing_content_types.write_bytes(b"PK\x03\x04not-a-zip")

    parser = DocxParser()

    assert parser.supports(wrong_extension) is False
    assert parser.supports(missing_content_types) is False


def test_pdf_support_requires_extension_and_magic(tmp_path: Path) -> None:
    parser = PdfParser(min_text_chars=0)
    pdf = tmp_path / "paper.pdf"
    pdf.write_bytes(b"%PDF-1.7\n%%EOF")
    disguised = tmp_path / "paper.txt"
    disguised.write_bytes(b"%PDF-1.7\n%%EOF")

    assert parser.supports(pdf) is True
    assert parser.supports(disguised) is False


@pytest.mark.parametrize(
    ("name", "parser"),
    [
        ("encrypted_pdf", PdfParser(min_text_chars=0)),
        ("javascript_pdf", PdfParser(min_text_chars=0)),
        ("giant_pdf", PdfParser(max_file_bytes=1024, min_text_chars=0)),
        ("macro_docx", DocxParser()),
        ("zip_bomb_docx", DocxParser(max_uncompressed_bytes=1024)),
        ("corrupted_docx", DocxParser()),
    ],
)
def test_malicious_documents_are_rejected(
    malicious_document_dir: dict[str, Path],
    name: str,
    parser: object,
) -> None:
    with pytest.raises(DocumentParseError):
        parser.parse(malicious_document_dir[name])  # type: ignore[attr-defined]
