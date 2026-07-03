import json
import os
import signal
import sys
import tempfile
import time
from concurrent.futures import ThreadPoolExecutor
from datetime import UTC, datetime
from pathlib import Path
from types import SimpleNamespace

import pytest
from docx import Document

import adapters.parser._sandbox as sandbox
from adapters.parser._sandbox import SANDBOX_TEMP_PREFIX, run_in_sandbox
from adapters.parser.docx_parser import DocxParser
from core.domain.exceptions import DocumentParseError
from core.interfaces.document_parser import (
    DocumentParser,
    ParsedDocument,
    ParsedLocatorIndex,
    compute_file_hash,
)

LARGE_PAYLOAD_CHARS = 80_000
VERY_LARGE_PAYLOAD_CHARS = 1_000_000


class ContextProbeParser(DocumentParser):
    def supports(self, file_path: Path) -> bool:
        return True

    def parse(self, file_path: Path, timeout_seconds: float = 30.0) -> ParsedDocument:
        _ = timeout_seconds
        payload = {
            "path": str(file_path),
            "cwd": os.getcwd(),
            "has_project_env": "MXA_PROJECT_ROOT" in os.environ,
        }
        return _parsed(file_path, json.dumps(payload))


class RaisingParser(DocumentParser):
    def supports(self, file_path: Path) -> bool:
        return True

    def parse(self, file_path: Path, timeout_seconds: float = 30.0) -> ParsedDocument:
        _ = timeout_seconds
        raise RuntimeError(f"boom {file_path} C:\\private\\secret.pdf /home/private")


class LargeRaisingParser(DocumentParser):
    def supports(self, file_path: Path) -> bool:
        return True

    def parse(self, file_path: Path, timeout_seconds: float = 30.0) -> ParsedDocument:
        _ = timeout_seconds, "x" * LARGE_PAYLOAD_CHARS
        raise RuntimeError(f"boom {file_path} C:\\private\\secret.pdf /home/private")


class PayloadParser(DocumentParser):
    def __init__(self, raw_text: str) -> None:
        self._raw_text = raw_text

    def supports(self, file_path: Path) -> bool:
        return True

    def parse(self, file_path: Path, timeout_seconds: float = 30.0) -> ParsedDocument:
        _ = timeout_seconds
        return _parsed(file_path, self._raw_text)


class SleepingParser(DocumentParser):
    def supports(self, file_path: Path) -> bool:
        return True

    def parse(self, file_path: Path, timeout_seconds: float = 30.0) -> ParsedDocument:
        time.sleep(timeout_seconds + 5)
        return _parsed(file_path, "unreachable")


class ExitParser(DocumentParser):
    def supports(self, file_path: Path) -> bool:
        return True

    def parse(self, file_path: Path, timeout_seconds: float = 30.0) -> ParsedDocument:
        _ = file_path, timeout_seconds
        os._exit(7)


class OutsidePathParser(DocumentParser):
    def supports(self, file_path: Path) -> bool:
        return True

    def parse(self, file_path: Path, timeout_seconds: float = 30.0) -> ParsedDocument:
        _ = file_path, timeout_seconds
        raise PermissionError("blocked outside path C:\\workspace\\project")


class CpuSpinParser(DocumentParser):
    def supports(self, file_path: Path) -> bool:
        return True

    def parse(self, file_path: Path, timeout_seconds: float = 30.0) -> ParsedDocument:
        _ = file_path, timeout_seconds
        counter = 0
        while True:
            counter = (counter + 1) % 1_000_000


def _put_wrong_payload_child(request: object, result_queue: object) -> None:
    _ = request
    result_queue.put(("ok", "not a parsed document"))  # type: ignore[attr-defined]


def _put_unknown_status_child(request: object, result_queue: object) -> None:
    _ = request
    result_queue.put(("unexpected", "document_parse_failed"))  # type: ignore[attr-defined]


def _put_bad_shape_child(request: object, result_queue: object) -> None:
    _ = request
    result_queue.put(("ok",))  # type: ignore[attr-defined]


def test_sandbox_child_receives_only_temp_path_and_config(tmp_path: Path) -> None:
    source = tmp_path / "secret_original_name.pdf"
    source.write_bytes(b"%PDF-1.7\n")

    parsed = run_in_sandbox(ContextProbeParser(), source)
    payload = json.loads(parsed.raw_text)

    assert "secret_original_name" not in payload["path"]
    assert payload["path"].endswith("document.pdf")
    assert str(Path.cwd()) not in payload["path"]


def test_sandbox_child_cwd_is_isolated_temp_dir(tmp_path: Path) -> None:
    source = tmp_path / "paper.pdf"
    source.write_bytes(b"%PDF-1.7\n")

    parsed = run_in_sandbox(ContextProbeParser(), source)
    payload = json.loads(parsed.raw_text)

    assert SANDBOX_TEMP_PREFIX in payload["cwd"]
    assert payload["cwd"] != str(Path.cwd())


def test_sandbox_child_env_does_not_include_project_root(
    tmp_path: Path,
    monkeypatch: pytest.MonkeyPatch,
) -> None:
    source = tmp_path / "paper.pdf"
    source.write_bytes(b"%PDF-1.7\n")
    monkeypatch.setenv("MXA_PROJECT_ROOT", str(Path.cwd()))

    parsed = run_in_sandbox(ContextProbeParser(), source)
    payload = json.loads(parsed.raw_text)

    assert payload["has_project_env"] is False


def test_parse_error_sanitizes_absolute_path_and_original_filename(tmp_path: Path) -> None:
    source = tmp_path / "private_report_name.pdf"
    source.write_bytes(b"%PDF-1.7\n")

    with pytest.raises(DocumentParseError) as exc_info:
        run_in_sandbox(RaisingParser(), source)

    message = str(exc_info.value)
    assert "private_report_name" not in message
    assert "C:\\" not in message
    assert "/home" not in message
    assert "/Users" not in message


def test_sandbox_timeout_is_reported_without_child_traceback(tmp_path: Path) -> None:
    source = tmp_path / "paper.pdf"
    source.write_bytes(b"%PDF-1.7\n")

    with pytest.raises(DocumentParseError) as exc_info:
        run_in_sandbox(SleepingParser(), source, timeout_seconds=0.2)

    assert str(exc_info.value) == "document_parse_timeout"


def test_sandbox_child_crash_is_isolated(tmp_path: Path) -> None:
    source = tmp_path / "paper.pdf"
    source.write_bytes(b"%PDF-1.7\n")

    with pytest.raises(DocumentParseError) as exc_info:
        run_in_sandbox(ExitParser(), source)

    assert str(exc_info.value) == "document_parse_failed"


def test_sandbox_error_if_parser_attempts_path_outside_temp_dir(tmp_path: Path) -> None:
    source = tmp_path / "paper.pdf"
    source.write_bytes(b"%PDF-1.7\n")

    with pytest.raises(DocumentParseError) as exc_info:
        run_in_sandbox(OutsidePathParser(), source)

    assert str(exc_info.value) == "document_parse_failed"


def test_sandbox_non_linux_can_fail_closed_when_hard_limits_required(tmp_path: Path) -> None:
    if sys.platform.startswith("linux"):
        pytest.skip("Linux supports hard parser limits")
    source = tmp_path / "paper.pdf"
    source.write_bytes(b"%PDF-1.7\n")

    with pytest.raises(DocumentParseError) as exc_info:
        run_in_sandbox(ContextProbeParser(), source, require_hard_limits=True)

    assert str(exc_info.value) == "unsupported_parser_sandbox_platform"


def test_sandbox_large_payload_round_trip_is_complete(tmp_path: Path) -> None:
    source = tmp_path / "paper.pdf"
    source.write_bytes(b"%PDF-1.7\n")
    raw_text = "x" * LARGE_PAYLOAD_CHARS

    parsed = run_in_sandbox(PayloadParser(raw_text), source, timeout_seconds=5.0)

    assert parsed.raw_text == raw_text


def test_sandbox_large_payload_parse_error_stays_sanitized(tmp_path: Path) -> None:
    source = tmp_path / "private_report_name.pdf"
    source.write_bytes(b"%PDF-1.7\n")

    with pytest.raises(DocumentParseError) as caught:
        run_in_sandbox(LargeRaisingParser(), source, timeout_seconds=5.0)

    message = caught.value.args[0]
    assert message == "document_parse_failed"
    assert "private_report_name" not in message
    assert "C:\\" not in message
    assert "/home" not in message
    assert "/Users" not in message


def test_sandbox_docx_parser_large_payload_round_trip(tmp_path: Path) -> None:
    source = tmp_path / "paper.docx"
    document = Document()
    for _ in range(20):
        document.add_paragraph("x" * 5_000)
    document.save(source)

    parsed = run_in_sandbox(DocxParser(), source, timeout_seconds=10.0)

    assert len(parsed.raw_text) >= LARGE_PAYLOAD_CHARS
    assert parsed.raw_text.count("x") == 100_000


def test_sandbox_child_crash_fails_fast(tmp_path: Path) -> None:
    source = tmp_path / "paper.pdf"
    source.write_bytes(b"%PDF-1.7\n")

    started = time.perf_counter()
    with pytest.raises(DocumentParseError) as caught:
        run_in_sandbox(ExitParser(), source, timeout_seconds=5.0)
    elapsed = time.perf_counter() - started

    assert caught.value.args == ("document_parse_failed",)
    assert elapsed < 2.0


def test_sandbox_timeout_cleans_process_and_temp_dir(tmp_path: Path) -> None:
    source = tmp_path / "paper.pdf"
    source.write_bytes(b"%PDF-1.7\n")
    temp_root = Path(tempfile.gettempdir())
    before = {path.name for path in temp_root.glob(f"{SANDBOX_TEMP_PREFIX}*") if path.is_dir()}

    with pytest.raises(DocumentParseError) as caught:
        run_in_sandbox(SleepingParser(), source, timeout_seconds=0.2)

    time.sleep(0.2)
    after = {path.name for path in temp_root.glob(f"{SANDBOX_TEMP_PREFIX}*") if path.is_dir()}
    assert caught.value.args == ("document_parse_timeout",)
    assert sorted(after - before) == []


def test_sandbox_concurrent_large_payload_calls_do_not_interfere(tmp_path: Path) -> None:
    sources: list[Path] = []
    for index in range(6):
        source = tmp_path / f"paper_{index}.pdf"
        source.write_bytes(b"%PDF-1.7\n")
        sources.append(source)

    def parse_one(source: Path) -> str:
        parsed = run_in_sandbox(
            PayloadParser(f"{source.stem}:" + "x" * LARGE_PAYLOAD_CHARS),
            source,
            timeout_seconds=5.0,
        )
        return parsed.raw_text

    with ThreadPoolExecutor(max_workers=3) as pool:
        results = list(pool.map(parse_one, sources))

    assert len(results) == len(sources)
    assert {result.split(":", 1)[0] for result in results} == {source.stem for source in sources}
    assert all(len(result.split(":", 1)[1]) == LARGE_PAYLOAD_CHARS for result in results)


def test_sandbox_dead_after_put_result_remains_visible(
    tmp_path: Path,
    monkeypatch: pytest.MonkeyPatch,
) -> None:
    source = tmp_path / "paper.pdf"
    source.write_bytes(b"%PDF-1.7\n")
    monkeypatch.setattr(sandbox, "POLL_SECONDS", 0.001)

    for size in (1, LARGE_PAYLOAD_CHARS, VERY_LARGE_PAYLOAD_CHARS) * 2:
        raw_text = "x" * size
        parsed = run_in_sandbox(PayloadParser(raw_text), source, timeout_seconds=5.0)
        assert parsed.raw_text == raw_text


def test_sandbox_timeout_shorter_than_poll_interval_is_not_rounded_up(tmp_path: Path) -> None:
    source = tmp_path / "paper.pdf"
    source.write_bytes(b"%PDF-1.7\n")

    started = time.perf_counter()
    with pytest.raises(DocumentParseError) as caught:
        run_in_sandbox(SleepingParser(), source, timeout_seconds=0.05)
    elapsed = time.perf_counter() - started

    assert caught.value.args == ("document_parse_timeout",)
    assert elapsed < 1.0


@pytest.mark.parametrize(
    "child_main",
    [_put_wrong_payload_child, _put_unknown_status_child, _put_bad_shape_child],
)
def test_sandbox_rejects_malformed_child_results(
    tmp_path: Path,
    monkeypatch: pytest.MonkeyPatch,
    child_main: object,
) -> None:
    source = tmp_path / "paper.pdf"
    source.write_bytes(b"%PDF-1.7\n")
    monkeypatch.setattr(sandbox, "_sandbox_child_main", child_main)

    with pytest.raises(DocumentParseError) as caught:
        run_in_sandbox(ContextProbeParser(), source, timeout_seconds=5.0)

    assert caught.value.args == ("document_parse_failed",)


@pytest.mark.skipif(not sys.platform.startswith("linux"), reason="RLIMIT_CPU is Linux-only")
def test_sandbox_cpu_limit_exit_maps_to_timeout(tmp_path: Path) -> None:
    source = tmp_path / "paper.pdf"
    source.write_bytes(b"%PDF-1.7\n")
    request = sandbox._SandboxChildRequest(
        parser=CpuSpinParser(),
        file_path=source.resolve(),
        sandbox_dir=tmp_path.resolve(),
        timeout_seconds=1.0,
        mem_limit_bytes=sandbox.DEFAULT_MEM_LIMIT_BYTES,
    )

    started = time.perf_counter()
    with pytest.raises(DocumentParseError) as caught:
        sandbox._run_child(request, timeout_seconds=10.0)
    elapsed = time.perf_counter() - started

    assert caught.value.args == ("document_parse_timeout",)
    assert elapsed < 8.0


@pytest.mark.skipif(
    not sys.platform.startswith("linux"), reason="SIGKILL classification is Linux-only"
)
def test_sandbox_sigkill_near_deadline_maps_to_timeout() -> None:
    process = SimpleNamespace(exitcode=-int(signal.SIGKILL))
    deadline = time.monotonic() + (sandbox.SIGKILL_DEADLINE_GRACE_SECONDS / 2)

    with pytest.raises(DocumentParseError) as caught:
        sandbox._raise_for_missing_result(process, deadline)

    assert caught.value.args == ("document_parse_timeout",)


@pytest.mark.skipif(
    not sys.platform.startswith("linux"), reason="SIGKILL classification is Linux-only"
)
def test_sandbox_sigkill_before_deadline_stays_failed() -> None:
    process = SimpleNamespace(exitcode=-int(signal.SIGKILL))
    deadline = time.monotonic() + sandbox.SIGKILL_DEADLINE_GRACE_SECONDS + 1.0

    with pytest.raises(DocumentParseError) as caught:
        sandbox._raise_for_missing_result(process, deadline)

    assert caught.value.args == ("document_parse_failed",)


def _parsed(file_path: Path, raw_text: str) -> ParsedDocument:
    return ParsedDocument(
        raw_text=raw_text,
        page_count=None,
        figure_placeholders=[],
        table_placeholders=[],
        locator_index=ParsedLocatorIndex(section_ids=["S1"], equation_ids=[], figure_ids=[]),
        file_hash=compute_file_hash(file_path),
        extracted_at=datetime.now(UTC),
    )
